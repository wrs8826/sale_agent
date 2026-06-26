"""内置工具的 LangChain @tool 版本。

本模块是所有内置工具的**单一实现来源**：
  - builtin_mcp_server.py 从这里导入函数，再封装成 MCP 工具（飞书路径）
  - QA 图的 call_tools_node 直接使用 BUILTIN_TOOLS 列表（网页路径）

新增工具：在本文件加一个 @tool 函数，再把它加入 BUILTIN_TOOLS 即可。
"""
from __future__ import annotations

from pathlib import Path
from typing import List, Optional, Tuple

from langchain_core.tools import tool

from agent_service.logging_config import get_logger

log = get_logger(__name__)

# ── 文档文本提取（工具与 /ingest 共用的单一实现） ──────────────────────────────

# 可直接按 UTF-8 读取的纯文本扩展名
_TEXT_EXTS = {".txt", ".md", ".rst", ".html", ".htm", ".csv", ".log", ".json"}


# ── PDF 提取：PyMuPDF 取正文文本 + pdfplumber 取表格，按页内位置合并去重 ─────────

def _point_in_any_bbox(x: float, y: float, bboxes: List[Tuple[float, float, float, float]]) -> bool:
    """判断点 (x, y) 是否落在任一表格 bbox (x0, top, x1, bottom) 内。"""
    for bx0, btop, bx1, bbottom in bboxes:
        if bx0 <= x <= bx1 and btop <= y <= bbottom:
            return True
    return False


def _render_pdf_table(rows: list) -> str:
    """把 pdfplumber 的 table.extract() 结果渲染为管道分隔文本（与 .docx 表格口径一致）。"""
    out = []
    for row in rows or []:
        cells = [((c or "").strip().replace("\n", " ")) for c in row]
        if any(cells):
            out.append(" | ".join(cells))
    return "\n".join(out)


def _extract_pdf(path: Path) -> str:
    """PDF 文本提取：PyMuPDF 取正文文本块 + pdfplumber 取表格，按页内垂直位置合并。

    - 落在 pdfplumber 识别的表格 bbox 内的 PyMuPDF 文本块会被剔除，避免与表格内容重复
      （PyMuPDF 全页文本本就包含被打散的表格单元格文字）。
    - pdfplumber 缺失或某页解析异常时，优雅降级为 PyMuPDF 全页文本，不报错。
    - 扫描版/图片型 PDF 仍无文本层（无 OCR），返回空由调用方提示。
    """
    import fitz  # PyMuPDF；调用方已确保依赖存在

    try:
        import pdfplumber
    except ImportError:
        log.warning("未安装 pdfplumber，PDF 表格将按普通文本提取；建议 pip install pdfplumber")
        pdfplumber = None

    if pdfplumber is None:
        with fitz.open(str(path)) as doc:
            return "\n".join(page.get_text() for page in doc).strip()

    page_texts: List[str] = []
    with fitz.open(str(path)) as fdoc, pdfplumber.open(str(path)) as pdoc:
        n = min(fdoc.page_count, len(pdoc.pages))
        for i in range(n):
            fpage = fdoc[i]
            ppage = pdoc.pages[i]

            # 1) pdfplumber 识别表格（bbox + 数据）；失败则该页退化为纯文本
            try:
                tables = ppage.find_tables() or []
            except Exception as exc:
                log.debug("pdfplumber 第 %d 页表格识别失败，退化为纯文本: %s", i + 1, exc)
                page_texts.append(fpage.get_text().strip())
                continue

            table_bboxes = [t.bbox for t in tables]  # (x0, top, x1, bottom)
            items: List[Tuple[float, float, str]] = []  # (top, left, content)

            # 2) PyMuPDF 文本块，剔除落在表格区域内的块（中心点判定）
            for blk in fpage.get_text("blocks"):
                # blk: (x0, y0, x1, y1, text, block_no, block_type)；block_type 1 为图片
                if len(blk) >= 7 and blk[6] != 0:
                    continue
                x0, y0, x1, y1, text = blk[0], blk[1], blk[2], blk[3], blk[4]
                if not text or not text.strip():
                    continue
                cx, cy = (x0 + x1) / 2, (y0 + y1) / 2
                if _point_in_any_bbox(cx, cy, table_bboxes):
                    continue
                items.append((y0, x0, text.strip()))

            # 3) 表格按其 bbox 顶部位置插回
            for t in tables:
                try:
                    rendered = _render_pdf_table(t.extract())
                except Exception as exc:
                    log.debug("pdfplumber 第 %d 页表格抽取失败: %s", i + 1, exc)
                    continue
                if rendered.strip():
                    items.append((t.bbox[1], t.bbox[0], rendered))

            # 4) 按垂直、再水平位置排序合并，复原阅读顺序
            items.sort(key=lambda it: (round(it[0], 1), it[1]))
            page_text = "\n".join(c for _, _, c in items).strip()
            if page_text:
                page_texts.append(page_text)

    return "\n".join(page_texts).strip()


# ── Word(.docx) 提取：unstructured 结构化为主，python-docx 为回退 ─────────────────

def _html_table_to_pipes(html: str) -> str:
    """把 unstructured Table 的 text_as_html 渲染为管道分隔文本（与 PDF/python-docx 表格口径一致）。"""
    if not html:
        return ""
    from html.parser import HTMLParser

    class _TableParser(HTMLParser):
        def __init__(self):
            super().__init__()
            self.rows: List[List[str]] = []
            self._row: Optional[List[str]] = None
            self._cell: Optional[List[str]] = None

        def handle_starttag(self, tag, attrs):
            if tag == "tr":
                self._row = []
            elif tag in ("td", "th"):
                self._cell = []

        def handle_endtag(self, tag):
            if tag == "tr" and self._row is not None:
                self.rows.append(self._row)
                self._row = None
            elif tag in ("td", "th") and self._cell is not None:
                if self._row is not None:
                    self._row.append("".join(self._cell).strip())
                self._cell = None

        def handle_data(self, data):
            if self._cell is not None:
                self._cell.append(data)

    parser = _TableParser()
    parser.feed(html)
    lines = []
    for row in parser.rows:
        cells = [c.replace("\n", " ").strip() for c in row]
        if any(cells):
            lines.append(" | ".join(cells))
    return "\n".join(lines)


def _extract_docx_unstructured(path: Path) -> Optional[str]:
    """用 unstructured 结构化读取 .docx：标题→Markdown 标题、表格→管道行、列表→`- `、其余→段落。

    元素按文档原始顺序返回（含页眉/页脚），标题渲染成 `## ` 便于下游 DocumentChunker 的标题分隔识别。
    返回结构化纯文本；unstructured 不可用或解析失败时返回 None，由调用方回退 python-docx。
    """
    try:
        from unstructured.partition.docx import partition_docx
    except ImportError:
        log.warning("未安装 unstructured，.docx 改用 python-docx 提取；建议 pip install unstructured")
        return None
    try:
        elements = partition_docx(filename=str(path))
    except Exception as exc:
        log.warning(".docx unstructured 解析失败，回退 python-docx: %s", exc)
        return None

    parts: List[str] = []
    for el in elements:
        category = getattr(el, "category", "") or type(el).__name__
        text = (getattr(el, "text", "") or "").strip()
        if category == "Table":
            html = getattr(getattr(el, "metadata", None), "text_as_html", None)
            rendered = _html_table_to_pipes(html) if html else text
            if rendered.strip():
                parts.append(rendered)
        elif category in ("Title", "Header"):
            if text:
                parts.append(f"## {text}")
        elif category == "ListItem":
            if text:
                parts.append(f"- {text}")
        else:
            if text:
                parts.append(text)
    return "\n".join(parts).strip()


def _extract_docx_python_docx(path: Path) -> str:
    """回退实现：python-docx 按 body 的 XML 子节点顺序遍历，保持段落与表格的原始先后位置。"""
    try:
        from docx import Document
        from docx.oxml.ns import qn
        from docx.table import Table
        from docx.text.paragraph import Paragraph
    except ImportError as exc:
        raise RuntimeError(
            "服务器未安装 python-docx，无法读取 Word，请先 pip install python-docx。"
        ) from exc
    doc = Document(str(path))
    # doc.paragraphs / doc.tables 是两个独立集合，分别遍历会把表格全部挪到文末；这里按 XML 顺序遍历。
    parts = []
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            text = Paragraph(child, doc).text
            if text and text.strip():
                parts.append(text.strip())
        elif child.tag == qn("w:tbl"):
            for row in Table(child, doc).rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def extract_text_from_file(path: Path) -> str:
    """按扩展名从文件中提取纯文本，供 read_document 工具与 /ingest 清洗共用。

    支持：
        .pdf            → PyMuPDF 取正文 + pdfplumber 取表格，按位置合并去重
        .docx           → unstructured 结构化提取（标题/表格/列表/页眉页脚）；失败回退 python-docx
        纯文本类         → 自动识别编码读取（UTF-8/GBK 等，见 text_utils.read_text_smart）
    不支持 .doc（旧二进制 Word）；缺少依赖或解析失败时抛出异常，由调用方决定如何呈现。

    Args:
        path: 目标文件的绝对路径（Path 对象）。

    Returns:
        提取出的纯文本（已 strip），不做长度截断。
    """
    ext = path.suffix.lower()

    if ext == ".pdf":
        try:
            import fitz  # noqa: F401  仅探测 PyMuPDF 依赖；实际提取在 _extract_pdf 内
        except ImportError as exc:
            raise RuntimeError(
                "服务器未安装 PyMuPDF，无法读取 PDF，请先 pip install pymupdf。"
            ) from exc
        return _extract_pdf(path)

    if ext == ".docx":
        # 主路径：unstructured 结构化读取；不可用或失败时回退 python-docx（按 body XML 顺序遍历）。
        text = _extract_docx_unstructured(path)
        if text is not None:
            return text
        return _extract_docx_python_docx(path)

    if ext == ".doc":
        raise RuntimeError("暂不支持旧版 .doc 二进制格式，请另存为 .docx 或 PDF 后再上传。")

    # 其余按纯文本读取（自动识别 UTF-8/GBK 等编码，避免中文乱码）
    from agent_service.text_utils import read_text_smart

    return read_text_smart(path)


# ── 工具定义 ──────────────────────────────────────────────────────────────────

@tool
def load_policy_file(skill_name: str, filename: str) -> str:
    """读取指定 skill 的政策文档文件，返回文件全文。

    适用场景：用户询问某个人才政策的具体细节（申报条件、资金政策、操作流程等）时，
    按**你系统提示里已给出的「文档地图」**选定目标文件后调用本工具读取原文。
    （文档地图已在系统提示中，无需再读 SKILL.md 获取；如确需，传 filename="SKILL.md" 也可读到。）

    Args:
        skill_name: skill 目录名，例如 "甬江人才政策"、"太仓人才政策"、
                    "无锡人才政策"、"成都人才政策"。
        filename:   文件名（含 .md 扩展名），例如 "申报条件_制造业.md"。
                    传文档地图里的文件名即可；带不带 "references/" 前缀都行（会自动处理）。
                    传 "SKILL.md" 可读取该 skill 的索引/文档地图本身。

    Returns:
        文件全文字符串；若文件不存在则返回错误说明（含可用文件清单）。
    """
    from pathlib import Path

    from agent_service import SKILLS_ROOT

    skill_dir = SKILLS_ROOT / skill_name
    safe = Path(filename).name  # 去掉任何路径前缀（如 references/）并防目录穿越
    # 先 references/，再 skill 根目录（SKILL.md 等索引文件在根目录）
    candidates = [skill_dir / "references" / safe, skill_dir / safe]
    path = next((p for p in candidates if p.is_file()), None)

    if path is None:
        refs_dir = skill_dir / "references"
        refs = [f.name for f in sorted(refs_dir.glob("*.md"))] if refs_dir.is_dir() else []
        roots = [f.name for f in sorted(skill_dir.glob("*.md"))] if skill_dir.is_dir() else []
        if refs or roots:
            return (f"文件 {filename!r} 不存在。{skill_name}/references/ 可用文件：{refs}；"
                    f"{skill_name}/ 根目录可用文件：{roots}")
        return f"文件 {filename!r} 不存在，或 skill {skill_name!r} 目录为空。"
    try:
        return path.read_text(encoding="utf-8")
    except Exception as exc:
        return f"读取文件失败：{exc}"


@tool
def get_current_time(timezone: str = "Asia/Shanghai") -> str:
    """获取当前日期和时间。

    Args:
        timezone: IANA 时区名称，例如 "Asia/Shanghai"（默认）、"UTC"、
                  "America/New_York"、"Europe/London"。

    Returns:
        格式为 "YYYY-MM-DD HH:MM:SS 时区缩写" 的字符串。
    """
    from datetime import datetime

    try:
        from zoneinfo import ZoneInfo, ZoneInfoNotFoundError
        try:
            tz = ZoneInfo(timezone)
        except (ZoneInfoNotFoundError, KeyError):
            tz = ZoneInfo("Asia/Shanghai")
    except ImportError:
        try:
            import pytz
            try:
                tz = pytz.timezone(timezone)
            except pytz.exceptions.UnknownTimeZoneError:
                tz = pytz.timezone("Asia/Shanghai")
        except ImportError:
            tz = None

    now = __import__("datetime").datetime.now(tz) if tz else __import__("datetime").datetime.now()
    return now.strftime("%Y-%m-%d %H:%M:%S %Z").strip()


def _is_table_row(line: str) -> bool:
    """判断是否为 Markdown 表格行（去空白后以 | 开头且含至少一个 |）。"""
    s = line.strip()
    return s.startswith("|") and s.count("|") >= 2


def _is_table_separator(line: str) -> bool:
    """判断是否为表格分隔行，如 |---|:--:|---| 。"""
    s = line.strip().strip("|")
    cells = [c.strip() for c in s.split("|")]
    return bool(cells) and all(c and set(c) <= set("-: ") for c in cells)


def _split_row(line: str) -> list:
    """切分一行表格为单元格列表（去掉首尾的 |）。"""
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.strip() for c in s.split("|")]


def _add_md_table(doc, rows: list) -> None:
    """把 Markdown 表格行（已剔除分隔行）渲染为带边框的 Word 表格，首行加粗作表头。"""
    cells_per_row = [_split_row(r) for r in rows]
    ncols = max(len(r) for r in cells_per_row)
    table = doc.add_table(rows=0, cols=ncols)
    try:
        table.style = "Table Grid"  # 带边框
    except Exception:
        pass
    for i, cells in enumerate(cells_per_row):
        row_cells = table.add_row().cells
        for j in range(ncols):
            text = cells[j] if j < len(cells) else ""
            row_cells[j].text = text
            if i == 0:  # 表头加粗
                for para in row_cells[j].paragraphs:
                    for run in para.runs:
                        run.bold = True


def _render_body(doc, body: str) -> None:
    """渲染章节正文：识别 Markdown 表格块转为 Word 表格，其余按段落输出。"""
    lines = body.split("\n")
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        if _is_table_row(line):
            # 收集连续的表格行
            block = []
            while i < n and _is_table_row(lines[i]):
                block.append(lines[i])
                i += 1
            # 过滤分隔行；若仍有内容则建表，否则退化为普通段落
            data_rows = [r for r in block if not _is_table_separator(r)]
            if data_rows:
                _add_md_table(doc, data_rows)
            continue
        text = line.strip()
        if text:
            doc.add_paragraph(text)
        i += 1


@tool
def generate_word_document(title: str, sections: list, filename: str = "") -> str:
    """把结构化内容生成为 Word(.docx) 文档并返回下载链接。

    适用场景：用户要求生成/导出软件说明书、设计文档、软著登记说明书、
    专利申请文件等可下载的 Word 文件时调用。先把要写入的内容组织成
    标题 + 章节列表，再调用本工具落盘。

    Args:
        title: 文档大标题，例如 "XX管理系统 软件说明书"。
        sections: 章节列表，每个元素是 dict：
            {
              "heading": "章节标题，如 一、概述",
              "body": "正文，可含多个段落（用 \\n 分隔）",
              "level": 1        # 可选，标题层级 1~3，默认 1，用于子章节缩进
            }
            正文留空（只给 heading）时仅输出标题，适合做分节占位。
        filename: 可选，自定义文件名（不含扩展名）；缺省用 title 生成。

    Returns:
        含 Markdown 下载链接的字符串；失败时返回错误说明。
        （网页端会自动渲染「下载」按钮，无需重复粘贴链接；飞书端由 MCP 转发层另行追加发链接指引。）
    """
    import re
    import uuid

    from agent_service import DOWNLOADS_DIR

    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        return "生成失败：服务器未安装 python-docx，请先 pip install python-docx。"

    if not title or not str(title).strip():
        return "生成失败：title 不能为空。"
    if not isinstance(sections, list) or not sections:
        return "生成失败：sections 必须是非空列表。"

    try:
        doc = Document()

        # 设置正文默认字体，确保中文正常显示（西文 + 东亚字体）
        normal = doc.styles["Normal"]
        normal.font.name = "宋体"
        normal.font.size = Pt(11)
        try:
            from docx.oxml.ns import qn
            normal.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        except Exception:
            pass

        doc.add_heading(str(title).strip(), level=0)

        rendered_sections = 0   # 实际写入了标题或正文的章节数；为 0 说明 sections 结构全无效
        for sec in sections:
            if not isinstance(sec, dict):
                continue
            heading = str(sec.get("heading") or "").strip()
            body = str(sec.get("body") or "").strip()
            try:
                level = int(sec.get("level", 1))
            except (TypeError, ValueError):
                level = 1
            level = min(max(level, 1), 4)

            if heading:
                doc.add_heading(heading, level=level)
            if body:
                _render_body(doc, body)
            if heading or body:
                rendered_sections += 1

        # 没有任何有效章节（如 LLM 误把 sections 传成 list[str]，dict 校验全跳过）：
        # 不能落一个只有标题的空文档却报「已生成」，否则用户下载到空文件。明确报错并提示正确结构。
        if rendered_sections == 0:
            return (
                "生成失败：sections 中没有有效章节。每个元素必须是包含 "
                '"heading" 和/或 "body" 的 dict，例如 '
                '[{"heading": "一、概述", "body": "正文…", "level": 1}]。'
            )

        # 文件名：清洗非法字符 + 短 uuid 防覆盖
        base = (filename or title).strip()
        base = re.sub(r'[\\/:*?"<>|]', "_", base)[:60] or "document"
        out_name = f"{base}_{uuid.uuid4().hex[:8]}.docx"

        DOWNLOADS_DIR.mkdir(parents=True, exist_ok=True)
        doc.save(str(DOWNLOADS_DIR / out_name))
    except Exception as exc:
        return f"生成 Word 文档失败：{exc}"

    # 仅陈述事实 + 给出真实链接，不命令模型粘贴：
    #   · 网页端由系统提示统一约定（系统自动渲染下载按钮、模型不要重复粘贴链接，见
    #     graph/qa/prompts._TOOLS_BODY）——若此处再写「请原样提供链接」会与系统提示冲突；
    #   · 飞书端无下载按钮 UI，由 builtin_mcp_server 转发层在重写为绝对链接后追加「请把链接发给用户」。
    # 链接本身（[..](/download/..docx)）保持不变，api 层 _extract_download 仍能抽出并下发 download 事件。
    return f"文档已生成：[{base}.docx](/download/{out_name})"


_READ_MAX_CHARS = 16000  # read_document 注入提示词的文本上限，超出截断


@tool
def list_documents() -> str:
    """列出知识库（docs）中当前可读取的所有文件名。

    适用场景：用户想知道知识库里有哪些文件，或在调用 read_document 读取整篇内容
    之前需要先确认准确文件名时调用。

    Returns:
        文件名列表字符串；知识库为空时给出提示。
    """
    from agent_service import DOCS_DIR

    if not DOCS_DIR.exists():
        return "知识库目录不存在。"
    files = sorted(f.name for f in DOCS_DIR.iterdir() if f.is_file())
    if not files:
        return "知识库当前为空。"
    return "知识库中的文件：\n" + "\n".join(f"- {n}" for n in files)


@tool
def read_document(filename: str) -> str:
    """读取知识库中某个文件的完整文本内容，支持 PDF / Word(.docx) / 纯文本文件。

    适用场景：用户要求阅读 / 总结 / 提取 / 针对某个【具体文件】问答，且需要的是
    整篇原文而非知识库检索片段时调用。文件名不确定时先调用 list_documents 确认。

    Args:
        filename: 文件名（含扩展名），例如 "合同.pdf"、"产品手册.docx"。
                  仅传文件名，不要带路径前缀。

    Returns:
        文件正文文本（过长时截断并附说明）；文件不存在时返回可用文件列表。
    """
    from agent_service import DOCS_DIR

    name = Path(filename).name  # 仅取文件名，防目录穿越
    target = DOCS_DIR / name
    if not target.is_file():
        available = (
            sorted(f.name for f in DOCS_DIR.iterdir() if f.is_file())
            if DOCS_DIR.exists() else []
        )
        if available:
            return f"文件 {name!r} 不存在。知识库中可用文件：{available}"
        return f"文件 {name!r} 不存在，且知识库为空。"

    try:
        text = extract_text_from_file(target)
    except Exception as exc:
        return f"读取文件失败：{exc}"

    if not text:
        return f"文件 {name!r} 未提取到文本内容（可能是扫描版 PDF 或空文件）。"

    if len(text) > _READ_MAX_CHARS:
        total = len(text)
        text = text[:_READ_MAX_CHARS] + (
            f"\n\n……（内容过长，仅展示前 {_READ_MAX_CHARS} 字，全文共 {total} 字）"
        )
    return f"【{name} 全文】\n{text}"


# ── 工具列表 ──────────────────────────────────────────────────────────────────
# BUILTIN_TOOLS：核心工具，网页端与飞书 QA 降级路径共用（飞书暂不含文档读取）。
# WEB_TOOLS：在核心基础上追加文档读取工具，仅用户端 / 管理员端（api/agent.py）启用。
BUILTIN_TOOLS = [get_current_time, load_policy_file, generate_word_document]
WEB_TOOLS = BUILTIN_TOOLS + [read_document, list_documents]


def build_tool_table(tools=None) -> str:
    """生成工具清单（Markdown 表格），供系统提示词常驻注入。

    取每个工具 docstring 的首行作为用途摘要；无工具时返回空串。

    Args:
        tools: 要展示的工具列表；缺省用 BUILTIN_TOOLS（核心集）。网页端应传入 WEB_TOOLS。
    """
    tools = tools if tools is not None else BUILTIN_TOOLS
    if not tools:
        return ""
    lines = ["| 工具 | 用途 |", "|---|---|"]
    for t in tools:
        desc = (t.description or "").strip()
        first = desc.splitlines()[0].strip() if desc else ""
        first = first.replace("|", "｜")  # 防止破坏表格
        lines.append(f"| {t.name} | {first} |")
    return "\n".join(lines)
